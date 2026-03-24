#!/usr/bin/env python3
"""產生教學用「公文函稿」空白 Word 範本至 ../assets/。需已安裝 python-docx（建議於 tools/.venv 內 pip install python-docx）。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    out = root / "assets" / "公文格式_空白範本_教學用.docx"

    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("臺北市政府勞動局函")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()
    doc.add_paragraph("（本檔為教學用空白範本；實際發文請依機關公文程式、用紙與簽核規定辦理。）")

    doc.add_paragraph("機關地址：【請填寫】")
    doc.add_paragraph("電　　話：【請填寫】　　傳　真：【請填寫】")
    doc.add_paragraph()

    doc.add_paragraph("受　文　者：【請填寫】")
    doc.add_paragraph("發文日期：中華民國　　　年　　　月　　　日")
    doc.add_paragraph("發文字號：【請填寫】")
    doc.add_paragraph("速　　別：【普通件／最速件等，請依實際】")
    doc.add_paragraph("密　　等及解密條件或保密期限：【請填寫；無則填無】")
    doc.add_paragraph("附　　件：【請填寫；如無附件請填「無」】")
    doc.add_paragraph()

    t = doc.add_paragraph()
    t.add_run("主旨：").bold = True
    t.add_run("【請填寫】")

    doc.add_paragraph()
    s = doc.add_paragraph()
    s.add_run("說　　明：").bold = True

    doc.add_paragraph("一、【請填寫】")
    doc.add_paragraph("二、【請填寫】")
    doc.add_paragraph("三、【請填寫】")
    doc.add_paragraph()

    doc.add_paragraph("正本：【請填寫】")
    doc.add_paragraph("副本：【請填寫】")
    doc.add_paragraph()

    doc.add_paragraph("承辦人：【請填寫】　　　　電話：【請填寫】")

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
